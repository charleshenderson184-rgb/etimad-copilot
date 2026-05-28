"""Universal document parser — extracts text from PDF, Word, Excel, images."""

from __future__ import annotations

import os
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from .pdf_processor import (
    PageContent,
    detect_document_language,
    detect_language,
    extract_text_from_pdf,
    get_full_text,
)

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pytesseract
    from PIL import Image
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".xlsx",
    ".xls",
    ".png",
    ".jpg",
    ".jpeg",
    ".txt",
}


@dataclass
class ParsedDocument:
    full_text: str
    language: str  # ar | en | mixed | unknown
    page_count: int
    file_type: str  # pdf | docx | xlsx | image | text
    pages: list[PageContent]


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".docx", ".doc"):
        return "docx"
    if ext in (".xlsx", ".xls"):
        return "xlsx"
    if ext in (".png", ".jpg", ".jpeg"):
        return "image"
    if ext == ".txt":
        return "text"
    return "unknown"


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def parse_docx(file_path: str) -> ParsedDocument:
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed")

    doc = DocxDocument(file_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)
    language = detect_language(full_text)

    page = PageContent(
        page_number=1,
        text=full_text,
        language=language,
        is_ocr=False,
    )

    return ParsedDocument(
        full_text=full_text,
        language=language if language in ("ar", "en") else detect_document_language([page]),
        page_count=1,
        file_type="docx",
        pages=[page],
    )


def parse_xlsx(file_path: str) -> ParsedDocument:
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl not installed")

    wb = load_workbook(file_path, data_only=True)
    sections: list[str] = []
    pages: list[PageContent] = []

    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        section_lines: list[str] = [f"=== Sheet: {sheet_name} ==="]

        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                section_lines.append(" | ".join(cells))

        sheet_text = "\n".join(section_lines)
        sections.append(sheet_text)
        pages.append(PageContent(
            page_number=sheet_idx + 1,
            text=sheet_text,
            language=detect_language(sheet_text),
            is_ocr=False,
        ))

    full_text = "\n\n".join(sections)

    return ParsedDocument(
        full_text=full_text,
        language=detect_document_language(pages),
        page_count=len(wb.sheetnames),
        file_type="xlsx",
        pages=pages,
    )


def parse_image(file_path: str) -> ParsedDocument:
    if not HAS_TESSERACT:
        raise RuntimeError(
            "Tesseract not installed. Install with: brew install tesseract tesseract-lang"
        )

    img = Image.open(file_path)
    text = pytesseract.image_to_string(img, lang="ara+eng").strip()
    language = detect_language(text)

    page = PageContent(
        page_number=1,
        text=text,
        language=language,
        is_ocr=True,
    )

    return ParsedDocument(
        full_text=text,
        language=language if language in ("ar", "en") else detect_document_language([page]),
        page_count=1,
        file_type="image",
        pages=[page],
    )


def parse_text(file_path: str) -> ParsedDocument:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    language = detect_language(text)
    page = PageContent(page_number=1, text=text, language=language, is_ocr=False)

    return ParsedDocument(
        full_text=text,
        language=language if language in ("ar", "en") else detect_document_language([page]),
        page_count=1,
        file_type="text",
        pages=[page],
    )


def parse_pdf(file_path: str) -> ParsedDocument:
    pages = extract_text_from_pdf(file_path)
    full_text = get_full_text(pages)
    language = detect_document_language(pages)

    return ParsedDocument(
        full_text=full_text,
        language=language,
        page_count=len(pages),
        file_type="pdf",
        pages=pages,
    )


def parse_document(file_path: str, filename: Optional[str] = None) -> ParsedDocument:
    """Universal entry point — dispatches to the right parser based on file extension."""
    name = filename or os.path.basename(file_path)
    file_type = get_file_type(name)

    if file_type == "pdf":
        return parse_pdf(file_path)
    if file_type == "docx":
        return parse_docx(file_path)
    if file_type == "xlsx":
        return parse_xlsx(file_path)
    if file_type == "image":
        return parse_image(file_path)
    if file_type == "text":
        return parse_text(file_path)

    raise ValueError(f"Unsupported file type: {Path(name).suffix}")
