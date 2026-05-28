"""Export generated proposals to Word (.docx) and PDF."""

from __future__ import annotations

import logging
import re
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)


def _is_arabic(text: str) -> bool:
    arabic_chars = sum(1 for c in text if "؀" <= c <= "ۿ" or "ݐ" <= c <= "ݿ")
    return arabic_chars > len(text) * 0.3


def _set_rtl(paragraph) -> None:
    """Apply right-to-left formatting to a paragraph (for Arabic)."""
    pPr = paragraph._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    if _is_arabic(text):
        _set_rtl(heading)
        for run in heading.runs:
            run.font.name = "Arial"
    else:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x06, 0x4E, 0x3B)


def _add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    if _is_arabic(text):
        _set_rtl(p)
        run.font.name = "Arial"


def _add_markdown_to_doc(doc: Document, markdown: str) -> None:
    """Render a markdown string into the Word document with reasonable formatting."""
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Skip JSON code fences (used for pricing data)
        if line.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                i += 1
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1)
        # Table detection
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            table_lines = [line]
            i += 2  # skip separator
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue
        # Bullet
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            if _is_arabic(line):
                _set_rtl(p)
        # Numbered
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(text, style="List Number")
            if _is_arabic(text):
                _set_rtl(p)
        # Regular paragraph
        else:
            # Handle inline bold
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
            _add_paragraph(doc, text)

        i += 1


def _add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Light Grid Accent 1"

    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            if col_idx >= col_count:
                break
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def render_proposal_to_docx(
    company_name: str,
    rfp_filename: str,
    executive_summary: Optional[str],
    technical_proposal: Optional[str],
    financial_proposal: Optional[str],
    language: str = "en",
) -> bytes:
    doc = Document()

    # Cover page
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = cover.add_run("TECHNICAL & FINANCIAL PROPOSAL")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x06, 0x4E, 0x3B)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("العرض الفني والمالي")
    sub_run.font.size = Pt(18)
    sub_run.font.name = "Arial"

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Submitted by: {company_name}\n").bold = True
    info.add_run(f"In response to: {rfp_filename}\n")

    doc.add_page_break()

    if executive_summary:
        _add_heading(doc, "Executive Summary / الملخص التنفيذي", level=1)
        _add_markdown_to_doc(doc, executive_summary)
        doc.add_page_break()

    if technical_proposal:
        _add_heading(doc, "Technical Proposal / العرض الفني", level=1)
        _add_markdown_to_doc(doc, technical_proposal)
        doc.add_page_break()

    if financial_proposal:
        _add_heading(doc, "Financial Proposal / العرض المالي", level=1)
        _add_markdown_to_doc(doc, financial_proposal)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def render_proposal_to_pdf(docx_bytes: bytes) -> Optional[bytes]:
    """Convert Word bytes to PDF via docx2pdf. Returns None if conversion isn't available."""
    try:
        import tempfile
        import os
        from docx2pdf import convert

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tmp_docx.write(docx_bytes)
            tmp_docx_path = tmp_docx.name

        tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
        convert(tmp_docx_path, tmp_pdf_path)

        with open(tmp_pdf_path, "rb") as f:
            pdf_bytes = f.read()

        os.unlink(tmp_docx_path)
        os.unlink(tmp_pdf_path)
        return pdf_bytes
    except Exception as e:
        logger.warning("PDF conversion failed: %s", e)
        return None
